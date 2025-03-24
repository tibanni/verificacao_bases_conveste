import pandas as pd
from datetime import datetime
import numpy as np
from docx import Document

#Baixando as bases e transformando em data frames
BaseAtual_Receber = pd.read_excel("JULHO_contas_receber.xlsx")
BaseAtual_Recebido = pd.read_excel("JULHO_contas_recebido.xlsx")
BaseAtual_produtos_gerais = pd.read_excel("JULHO_produtos_gerais.xlsx")
BaseAtual_vendas = pd.read_excel("JULHO_vendas.xlsx")
BaseAtualizacao_Receber = pd.read_excel("AGOSTO_contas_receber.xlsx")
BaseAtualizacao_Recebido = pd.read_excel("AGOSTO_contas_recebido.xlsx")
BaseAtualizacao_produtos_gerais = pd.read_excel("AGOSTO_produtos_gerais.xlsx")
BaseAtualizacao_vendas = pd.read_excel("AGOSTO_vendas.xlsx")

#Criando tabela Para armazenar as verificações
dados_verificacoes = []

#Função para verificar as mudanças de campos

def registros_mod(df_antigo, df_novo, chave):
    df_merged = pd.merge(df_antigo, df_novo, on=chave, suffixes=('_julho', '_agosto'), how='inner')
    modificados = []

    for index, row in df_merged.iterrows():
        mudancas_registro = {chave: row[chave]}  # Inicializa o dicionário para o registro atual
        teve_mudanca = False  # Flag para verificar se houve alguma mudança no registro

        for coluna in df_antigo.columns:
            if (coluna != chave) and (coluna != "DataReferencia") and (coluna != "DataReferência"):
                valor_antigo = row[coluna + '_julho']
                valor_novo = row[coluna + '_agosto']

                # Converte para string para evitar problemas de comparação de tipos
                valor_antigo = str(valor_antigo)
                valor_novo = str(valor_novo)

                if valor_antigo != valor_novo:
                    mudancas_registro[coluna] = {'julho': valor_antigo, 'agosto': valor_novo}
                    teve_mudanca = True
                else:
                    mudancas_registro[coluna] = ""  # ou "" para deixar vazio

        if teve_mudanca:
            modificados.append(mudancas_registro)

    if modificados:
        df_modificados = pd.DataFrame(modificados)
        return df_modificados
    else:
        print("\nNenhum registro modificado.")
        return pd.DataFrame()

#Funcao que gera um relatório em formato docx com base nos dados fornecidos.
#Args:dados: Uma lista de listas, onde cada lista interna representa uma verificação com os seguintes elementos: base, verificação, situação.

def gerar_relatorio(dados):
  
  document = Document()
  document.add_heading('Verificações nas Bases', 0)

  bases = {}
  for base, verificacao, situacao in dados:
      if base not in bases:
          bases[base] = []
      bases[base].append((verificacao, situacao))

  for base, verificacoes in bases.items():
      document.add_heading(base, level=1)
      for verificacao, situacao in verificacoes:
          p = document.add_paragraph(verificacao)
          p.add_run(' - ').bold = True

          if situacao:
            p.add_run("Conforme")
          else:
            p.add_run("Não conforme")

  document.save('relatorio_verificacoes.docx')

#..............................................................................................................................................................................................
#...............................................................................VERIFICAÇÃO BASE PRODUTOS GERAIS...............................................................................
#..............................................................................................................................................................................................

#Criando coluna nas bases com a concatenação da obra e identificador da unidade
BaseAtual_produtos_gerais["Obra_IdentUnid"] = BaseAtual_produtos_gerais["Obra"].astype(str)+BaseAtual_produtos_gerais["Identificador"].astype(str)
BaseAtualizacao_produtos_gerais["Obra_IdentUnid"] = BaseAtualizacao_produtos_gerais["Obra"].astype(str)+BaseAtualizacao_produtos_gerais["Identificador"].astype(str)
BaseAtualizacao_vendas["Obra_IdentUnid"] = BaseAtualizacao_vendas["descr_obr"].astype(str) + BaseAtualizacao_vendas["Identificador_unid"].astype(str)

#Vericação da correspondência da atual na atualização e verificando unicidade das unidades
Obra_Identificador_atual = BaseAtual_produtos_gerais["Obra_IdentUnid"].isin(BaseAtualizacao_produtos_gerais["Obra_IdentUnid"])
dados_verificacoes.append(["BASE PRODUTOS GERAIS", "Correspondencia da atual na atualizacao.", Obra_Identificador_atual.all()])

nao_corresp_id_atual = BaseAtual_produtos_gerais[~Obra_Identificador_atual]

#Verificação da atualização na atual
Obra_Identificador_atualizacao = BaseAtualizacao_produtos_gerais["Obra_IdentUnid"].isin(BaseAtual_produtos_gerais["Obra_IdentUnid"])
dados_verificacoes.append(["BASE PRODUTOS GERAIS", "Correspondencia da atualizacao na atual.", Obra_Identificador_atualizacao.all()])

nao_corresp_id_atualizacao = BaseAtualizacao_produtos_gerais[~Obra_Identificador_atualizacao]

#Verificação das Dublicatas
PD_unicos = BaseAtualizacao_produtos_gerais["Obra_IdentUnid"].is_unique
dados_verificacoes.append(["BASE PRODUTOS GERAIS", "Verificando unicidade das unidades.", PD_unicos])
#Caso tenha valores duplicados, esse são as unidades duplicadas
mascara = BaseAtualizacao_produtos_gerais["Obra_IdentUnid"].value_counts()>1
mascara_base_Pg = BaseAtualizacao_produtos_gerais["Obra_IdentUnid"].isin(mascara[mascara.values == True].index)
ObraUnid_Duplicada = BaseAtualizacao_produtos_gerais[mascara_base_Pg]

#Verificar Mudança de Status
#Funcao que recebe os Obra_ID que tiveram mudanca de status e compara eles, verificando se teve mudanca nao real.
def comparar_classificacao(dicionario):
    VIAtualizacao = dicionario['agosto']
    VIAtual = dicionario['julho']

    status = ["10-Dação", "8-Fora de Venda", "7-Suspenso p/ Venda"]
    status2 = ["0-Disponível", "1-Vendido", "4-Quitado"]
    
    if (VIAtual in status) & (VIAtualizacao in status2):
        minimud = "Analisar - Verificação 1"
        
    #Verifica as mudanças de status "4-Quitado" ou "1-Vendido" para "10-Dação" ou "8-Fora de Venda" ou "7-Suspenso p/ Venda"
    elif (VIAtual in status2) & (VIAtualizacao in status):
        minimud = "Analisar - Verificação 2"
    
    else:
        minimud = "Ok"
        
    return minimud

#Verifica as modificacoes de campos
PG_mod_campos = registros_mod(BaseAtual_produtos_gerais, BaseAtualizacao_produtos_gerais, "Obra_IdentUnid")
#Filtrado registro que tiveram mudanca na 'Classificacao'
mudancas_status = PG_mod_campos[['Obra_IdentUnid', 'Classificação']][PG_mod_campos["Classificação"] != ""].copy()
#Aplicando funcao para vareificar mudanca de estado
mudancas_status["Conforme"] = mudancas_status["Classificação"].apply(comparar_classificacao)

#Verificacao de status "Vendida" ou "Quitada" para correspondencia na base de vendas
unidades_vendas = BaseAtualizacao_vendas[(BaseAtualizacao_vendas["Status Venda"] == "QUITADA") | (BaseAtualizacao_vendas["Status Venda"] == "NORMAL")][["Obra_IdentUnid","Status Venda","vendaId"]].copy()
unidades_vendidas = BaseAtualizacao_produtos_gerais[(BaseAtualizacao_produtos_gerais["Classificação"] == "4-Quitado") | (BaseAtualizacao_produtos_gerais["Classificação"] == "1-Vendido")][["Obra_IdentUnid","Obra","Identificador","Classificação","Nº Venda"]].copy()
comparativo_vendas_id = pd.merge(unidades_vendas, unidades_vendidas, on="Obra_IdentUnid", suffixes=("_vendas", "_pg"))
comparativo_vendas_id["Compara_VendaID"] = comparativo_vendas_id['vendaId'] == comparativo_vendas_id["Nº Venda"]

#Comparativo nas mudancas de status
#Contabiliza quantas mudanças de classificacao teve, a partir da diferença da quantidade de cada nas tabelas atual e de atualizacao. Qtd(Atualizacao) - Qtd(Atual)
classificacoes_atual = BaseAtual_produtos_gerais.groupby("Classificação")["Classificação"].count()
classificacoes_atualizada = BaseAtualizacao_produtos_gerais.groupby("Classificação")["Classificação"].count()
itens_unicos = pd.Index(classificacoes_atual.index.union(classificacoes_atualizada.index))
classificacoes_atual = classificacoes_atual.reindex(itens_unicos, fill_value=0)
classificacoes_atualizada = classificacoes_atualizada.reindex(itens_unicos, fill_value=0)
diferenca_classificacao = classificacoes_atualizada - classificacoes_atual

#Contabiliza quantas mudanças de classificação teve, a partir da contabilizacao das mudancas que teve na tabela de modificacoes. Para uma Classificacao no mes atualizacao e + e atual e -
contagem = mudancas_status['Classificação'].apply(lambda x: (x['julho'], x['agosto']))
contagem_mudanca_status = contagem.value_counts()

#Faz a contabilizacao da diferenca pela modificacao de Classificacao
resumo_status = {}
for i, n in contagem_mudanca_status.items():
    
    if i[1] not in resumo_status:
        resumo_status[i[1]] = n
        
    else:
        resumo_status[i[1]] = resumo_status.get(i[1]) + n
        
    if i[0] not in resumo_status:
        resumo_status[i[0]] = -n
        
    else:
        resumo_status[i[0]] = resumo_status.get(i[0]) - n

#Compara o valor das diferenças
classificacao_incorreta = {"Classificacao_Incorreta":[], "Dif_Bases":[], "Dif_Mud":[]}
for j, k in diferenca_classificacao.items():
    if j in resumo_status:
        if k != resumo_status[j]:
            classificacao_incorreta["Classificacao_Incorreta"].append(j)
            classificacao_incorreta["Dif_Bases"].append(k)
            classificacao_incorreta["Dif_Mud"].append(resumo_status[j])
        else:
            b = 1+1

    else:
        if k > 0:
            classificacao_incorreta["Classificacao_Incorreta"].append(j)
            classificacao_incorreta["Dif_Bases"].append(k)
            classificacao_incorreta["Dif_Mud"].append(0) 
        else:
            a = 1+1

df_classificacao_incorreta = pd.DataFrame(classificacao_incorreta)

#Mudancas de valor da area das unidades
mudancas_qtd_m2 = PG_mod_campos[['Obra_IdentUnid', 'Qtde M²']][PG_mod_campos["Qtde M²"] != ""].copy()

#Mudanças de Valor da unidade. Desconsiderando quem mudou de Classificação
#Funcao para acessar os valores de mudanca e comparar a variacao
def comparar_valor(dicionario):
    valor_atualizacao = float(dicionario['agosto'])
    valor_atual = float(dicionario['julho'])

    variacao = (abs(valor_atual - valor_atualizacao)/valor_atual)

    if variacao >= 0.25:
        minimud = "Analisar Varicação"     

    else:
        minimud = "Ok"         

    return minimud

mudancas_valor = PG_mod_campos[['Obra_IdentUnid','Valor']][(PG_mod_campos["Classificação"] == "") & (PG_mod_campos["Valor"] != "")].copy()
mudancas_valor["Variação"] = mudancas_valor["Valor"].apply(comparar_valor)

PG_mod_campos.to_excel("PG_mod_campos.xlsx", index=False)

with pd.ExcelWriter("PG_verificacoes.xlsx") as writer:
    nao_corresp_id_atual.to_excel(writer,sheet_name="PG_nao_corresp_atual", index=False)
    nao_corresp_id_atualizacao.to_excel(writer,sheet_name="PG_nao_corresp_atualizacao", index=False)
    ObraUnid_Duplicada.to_excel(writer,sheet_name="PG_ObraUnid_Duplicada", index=False)
    mudancas_status.to_excel(writer,sheet_name="Mudancas_Status_PG", index=False)
    comparativo_vendas_id.to_excel(writer,sheet_name="Conferencia_VendaID", index=False)
    df_classificacao_incorreta.to_excel(writer,sheet_name="Conferencia_Dif_Classificacao", index=False)
    mudancas_qtd_m2.to_excel(writer,sheet_name="Mudancas_Qtd_M2", index=False)
    mudancas_valor.to_excel(writer,sheet_name="Conferencia_Mudancas_Valor", index=False)


#..............................................................................................................................................................................................
#.............................................................................VERIFICAÇÃO VENDAS...............................................................................................
#..............................................................................................................................................................................................

#Modificacoes que teve
VENDAS_mod_campos = registros_mod(BaseAtual_vendas, BaseAtualizacao_vendas, "vendaId")

BaseAtualizacao_vendas["Obra_IdentUnid"] = BaseAtualizacao_vendas["descr_obr"].astype(str) + BaseAtualizacao_vendas["Identificador_unid"].astype(str)

#Verificação das Dublicatas
VENDAS_unicos = BaseAtualizacao_vendas["vendaId"].is_unique
dados_verificacoes.append(["BASE VENDAS", "Verificando unicidade das VENDAS.", VENDAS_unicos])
#Caso tenha valores duplicados, esse são as unidades duplicadas    
mascara_vendaID_Duplciada = BaseAtualizacao_vendas["vendaId"].value_counts() > 1
mascara_vendas_duplicadas = BaseAtualizacao_vendas["vendaId"].isin(mascara_vendaID_Duplciada[mascara_vendaID_Duplciada.values == True].index)
vendas_duplicadas = BaseAtualizacao_vendas[mascara_vendas_duplicadas]

# Verificar se as vendas na base atual constam na de atualização. 
VendaID_atual = BaseAtual_vendas["vendaId"].isin(BaseAtualizacao_vendas["vendaId"])
dados_verificacoes.append(["BASE VENDAS", "Verifica se Vendas da atual constam na Atualização.", VendaID_atual.all()])
nao_corresp_vendas_atual = BaseAtual_vendas[~VendaID_atual]

#Verificação da atualização na atual
VendaID_atualizacao = BaseAtualizacao_vendas["vendaId"].isin(BaseAtual_vendas["vendaId"])
nao_corresp_vendas_atualizacao = BaseAtualizacao_vendas[~VendaID_atualizacao]
#verifica as vendas da base atualização, se são vendas do mês de referência
nao_corresp_vendas_atualizacao["data_ven"] = nao_corresp_vendas_atualizacao["data_ven"].dt.strftime('%Y-%m')
data_referencia = nao_corresp_vendas_atualizacao["DataReferencia"].dt.strftime('%Y-%m')
dt = data_referencia[data_referencia.index[0]]
tabela_analisar = nao_corresp_vendas_atualizacao[nao_corresp_vendas_atualizacao["data_ven"] != dt]
novas_vendas = nao_corresp_vendas_atualizacao[nao_corresp_vendas_atualizacao["data_ven"] == dt]

#Verificar para cada obra_ven+identificador_unid se tem apenas uma venda com o status “Normal” ou “Quitada”. Relação Venda_id para uma unidade com status normal ou quitada. 
Obra_Unid_unico = BaseAtualizacao_vendas[(BaseAtualizacao_vendas["Status Venda"] == "NORMAL") | (BaseAtualizacao_vendas["Status Venda"] == "QUITADA")]["Obra_IdentUnid"].is_unique
dados_verificacoes.append(["BASE VENDAS", "Verificando unicidade das Obras_Unid por venda Normal ou Quitada.", Obra_Unid_unico])
  
mascara_obra_un = BaseAtualizacao_vendas[(BaseAtualizacao_vendas["Status Venda"] == "NORMAL") | (BaseAtualizacao_vendas["Status Venda"] == "QUITADA")]["Obra_IdentUnid"].value_counts() > 1
vendas_normal_quitada = BaseAtualizacao_vendas[(BaseAtualizacao_vendas["Status Venda"] == "NORMAL") | (BaseAtualizacao_vendas["Status Venda"] == "QUITADA")]
mascara_vendas_norm_quit = vendas_normal_quitada["Obra_IdentUnid"].isin(mascara_obra_un[mascara_obra_un.values == True].index)
obraID_duplicado_NormQuit = vendas_normal_quitada[mascara_vendas_norm_quit]
#
#Se todas as vendas com status quitada tem parcela na base de recebidas. Desconsiderar as vendas com data de cessão.
mascara_vendas_quitada_recebida = BaseAtualizacao_vendas[BaseAtualizacao_vendas["Status Venda"]=="QUITADA"]["vendaId"].isin(BaseAtualizacao_Recebido["vendaId"])
vendas_quitada = BaseAtualizacao_vendas[BaseAtualizacao_vendas["Status Venda"]=="QUITADA"]
vendas_sem_corresp = vendas_quitada[~mascara_vendas_quitada_recebida]
#Tratamento. Eliminando casos de cessão, com data de cessão e sem data de quitação
vendas_sem_corresp = vendas_sem_corresp[(vendas_sem_corresp["DataCessao_Ven"].isnull())]

if vendas_sem_corresp.shape[0] == 0:
    VD_Quit_Recebidas = True
else:
    VD_Quit_Recebidas = False
dados_verificacoes.append(["BASE VENDAS", "Verificando se tem recebimento das vendas com status QUITADA.", VD_Quit_Recebidas])

#Se todas as vendas com status “Normal” têm parcela na base a receber.
mascara_vendas_normal_recebida = BaseAtualizacao_vendas[BaseAtualizacao_vendas["Status Venda"]=="NORMAL"]["vendaId"].isin(BaseAtualizacao_Receber["vendaId"])
dados_verificacoes.append(["BASE VENDAS", "Verificando se tem parcelas a receber das vendas com status Normal.", mascara_vendas_normal_recebida.all()])
vendas_normal = BaseAtualizacao_vendas[BaseAtualizacao_vendas["Status Venda"]=="NORMAL"]
vendas_sem_corresp_normal = vendas_normal[~mascara_vendas_normal_recebida]
#
#Verificação da ficha financeira
ficha_financeira_vendas = BaseAtualizacao_vendas[(BaseAtualizacao_vendas["Status Venda"] == "NORMAL") | (BaseAtualizacao_vendas["Status Venda"] == "QUITADA")][["vendaId", "obra_ven","data_ven", "Status Venda","DataCessao_Ven", "valorTot_ven", "desconto_ven", "acrescimo_ven", "totVenda"]]
#Preenchendo com os valores a receber
ficha_financeira_vendas.Receber = 0.0

for v_id in ficha_financeira_vendas[ficha_financeira_vendas["Status Venda"] == "NORMAL"]["vendaId"]:
    valor_receber = BaseAtualizacao_Receber[BaseAtualizacao_Receber["vendaId"] == v_id]["ValParcela_crc"].sum()
    valor_receber = round(valor_receber, 2)

    ficha_financeira_vendas.loc[ficha_financeira_vendas["vendaId"] == v_id, "Receber"] =  valor_receber

#Preenchendo com os valores recebidos
ficha_financeira_vendas.Recebido = 0.0

for v_id in ficha_financeira_vendas["vendaId"]:
    valor_recebido = BaseAtualizacao_Recebido[BaseAtualizacao_Recebido["vendaId"] == v_id]["TotParcel"].sum()
    valor_recebido = round(valor_recebido, 2)

    ficha_financeira_vendas.loc[ficha_financeira_vendas["vendaId"] == v_id, "Recebido"] =  valor_recebido

ficha_financeira_vendas["Receber"].fillna(0, inplace = True)
ficha_financeira_vendas["Recebido"].fillna(0, inplace = True)    
#Fazendo a coluna total
ficha_financeira_vendas["Total_Parcelas"] =  ficha_financeira_vendas["Recebido"] + ficha_financeira_vendas["Receber"]
#Comparando valor de venda do produto com valor total das parcelas dele, receber+recebido. Desconsiderar Vendas com status quitada
ficha_financeira_vendas["Parcelas_maior_Vendas"] = ficha_financeira_vendas["totVenda"] <= ficha_financeira_vendas["Total_Parcelas"]

#Comparativo nas mudancas de status
#Contabiliza quantas mudanças de classificacao teve, a partir da diferença da quantidade de cada nas tabelas atual e de atualizacao. Qtd(Atualizacao) - Qtd(Atual)
status_atual = BaseAtual_vendas.groupby("Status Venda")["Status Venda"].count()
status_atualizada = BaseAtualizacao_vendas.groupby("Status Venda")["Status Venda"].count()
itens_unicos_vendas = pd.Index(status_atual.index.union(status_atualizada.index))
status_atual = status_atual.reindex(itens_unicos_vendas, fill_value=0)
status_atualizada = status_atualizada.reindex(itens_unicos_vendas, fill_value=0)
diferenca_status = status_atualizada - status_atual

#Contabiliza quantas mudanças de classificação teve, a partir da contabilizacao das mudancas que teve na tabela de modificacoes. Para uma Classificacao no mes atualizacao e + e atual e -
contagem = VENDAS_mod_campos['Status Venda'][VENDAS_mod_campos['Status Venda'] != ""].apply(lambda x: (x['julho'], x['agosto']))
contagem_mudanca_status_venda = contagem.value_counts()

#Faz a contabilizacao da diferenca pela modificacao de Classificacao
resumo_status = {}
for i, n in contagem_mudanca_status_venda.items():
    
    if i[1] not in resumo_status:
        resumo_status[i[1]] = n
        
    else:
        resumo_status[i[1]] = resumo_status.get(i[1]) + n
        
    if i[0] not in resumo_status:
        resumo_status[i[0]] = -n
        
    else:
        resumo_status[i[0]] = resumo_status.get(i[0]) - n

#Corrigindo a contagem pela diferenca com as novas vendas
status_atualizada_ref = novas_vendas.groupby("Status Venda")["Status Venda"].count()
for i, v in status_atualizada_ref.items():
    resumo_status[i] = resumo_status[i] + v

#Compara o valor das diferenças
status_incorreto = {"Classificacao_Incorreta":[], "Dif_Bases":[], "Dif_Mud":[]}
for j, k in diferenca_status.items():
    if j in resumo_status:
        if k != resumo_status[j]:
            status_incorreto["Classificacao_Incorreta"].append(j)
            status_incorreto["Dif_Bases"].append(k)
            status_incorreto["Dif_Mud"].append(resumo_status[j])
        else:
            b = 1+1

    else:
        if k > 0:
            status_incorreto["Classificacao_Incorreta"].append(j)
            status_incorreto["Dif_Bases"].append(k)
            status_incorreto["Dif_Mud"].append(0) 
        else:
            a = 1+1

df_status_incorreto = pd.DataFrame(status_incorreto)

#Mudanca valor Tot Venda
mud_valor_tot_venda = VENDAS_mod_campos[VENDAS_mod_campos['valorTot_ven'] != ""]

#Avalicao mudanca no valor totVenda. Verificar variacao maior que 25%. Funcao na parte de Produtos Gerais
mudancas_valor_venda = VENDAS_mod_campos[['vendaId','totVenda']][(VENDAS_mod_campos["totVenda"] != "")].copy()
mudancas_valor_venda["Variação"] = mudancas_valor_venda["totVenda"].apply(comparar_valor)

#Verificação outliers de desconto e valor de venda
def detect_outliers(df, date_col, value_col, threshold):
    window_size=365

    #Converter a coluna de data para datetime
    df[date_col] = pd.to_datetime(df[date_col])

    # Ordenar os dados por data
    df = df.sort_values(by=date_col)

    # Calcular a média móvel e o desvio padrão
    df['rolling_mean'] = df[value_col].rolling(window=window_size).mean()
    df['rolling_std'] = df[value_col].rolling(window=window_size).std()

    # Identificar outliers
    df['is_outlier'] = (abs(df[value_col] - df['rolling_mean']) > threshold * df['rolling_std'])

    return df

OBRAS = BaseAtual_vendas["obra_ven"].unique().tolist()
avaliar_out = ficha_financeira_vendas[["obra_ven","vendaId", "data_ven", "valorTot_ven","desconto_ven", "totVenda"]]
avaliar_out["Diff_Tot_TotVen"] = avaliar_out["totVenda"] - avaliar_out["valorTot_ven"] 

#Avaliando o Valor Total de venda por obre
out_vendas_total = pd.DataFrame()
out_vendas_total_2 = pd.DataFrame()
for obra in OBRAS:

    avaliar_out_tot = avaliar_out[avaliar_out["obra_ven"] ==  obra]
    if OBRAS.index(obra) == 1:
        
        out_vendas_total =  detect_outliers(avaliar_out_tot, "data_ven", "valorTot_ven", 7)
    
    else:
       
        out_vendas_total_2 =  detect_outliers(avaliar_out_tot, "data_ven", "valorTot_ven", 7)

    out_vendas_total = pd.concat([out_vendas_total, out_vendas_total_2], ignore_index=True)

#Avaliando o Valor Total de desconto por obra
out_vendas_desconto = pd.DataFrame()
out_vendas_desconto_2 = pd.DataFrame()
for obra in OBRAS:
    avaliar_out_desconto = avaliar_out[avaliar_out["obra_ven"] ==  obra]
    if OBRAS.index(obra) == 1:
        
        out_vendas_desconto =  detect_outliers(avaliar_out_desconto, "data_ven", "desconto_ven", 10)
    
    else:
        out_vendas_desconto_2 =  detect_outliers(avaliar_out_desconto, "data_ven", "desconto_ven", 10)

    out_vendas_desconto = pd.concat([out_vendas_desconto, out_vendas_desconto_2], ignore_index=True)

#Avaliando o Valor Total final da venda por obra
out_vendas_total_final = pd.DataFrame()
out_vendas_total_final_2 = pd.DataFrame()
for obra in OBRAS:
    avaliar_out_total_final = avaliar_out[avaliar_out["obra_ven"] ==  obra]
    if OBRAS.index(obra) == 1:
        
        out_vendas_total_final =  detect_outliers(avaliar_out_total_final, "data_ven", "Diff_Tot_TotVen", 10)
    
    else:
        out_vendas_total_final_2 =  detect_outliers(avaliar_out_total_final, "data_ven", "Diff_Tot_TotVen", 10)

    out_vendas_total_final = pd.concat([out_vendas_total_final, out_vendas_total_final_2], ignore_index=True)

#Vinculando vendas de cessao aos vendaId das vendas originarias dela
vendas_cessao = BaseAtualizacao_vendas[BaseAtualizacao_vendas['DataCessao_Ven'].notna()].copy()
vendas_cessao["VendaID_Cessao"] = ""

for row in vendas_cessao[['Obra_IdentUnid','data_ven', 'DataCessao_Ven', 'vendaId']].itertuples(index=False):
    obraid, dtVen, dtCes, vendaidCess = row
    df_filtrado = BaseAtualizacao_vendas[(BaseAtualizacao_vendas["Obra_IdentUnid"] == obraid) & (BaseAtualizacao_vendas["data_ven"] == dtVen) & (BaseAtualizacao_vendas["DataCancel_Ven"] == dtCes)]
    vendas_cessao["VendaID_Cessao"][vendas_cessao["vendaId"] == vendaidCess] = df_filtrado.iloc[0]['vendaId']

#Salvando as tabelas de verificacoes no Excel
with pd.ExcelWriter("VENDAS_verificacoes.xlsx") as writer:
    vendas_duplicadas.to_excel(writer, sheet_name="VENDAS_Vendas_Duplicadas", index=False)
    nao_corresp_vendas_atual.to_excel(writer, sheet_name="VENDAS_nao_corresp_vendas_atual", index=False)
    tabela_analisar.to_excel(writer, sheet_name="VENDAS_Novos_registros_inconsistente", index=False)
    obraID_duplicado_NormQuit.to_excel(writer, sheet_name="VENDAS_obraID_duplicado_NormQuit", index=False)
    vendas_sem_corresp.to_excel(writer, sheet_name="VENDAS_quitadas_sem_corresp_recebida", index=False)
    vendas_sem_corresp_normal.to_excel(writer, sheet_name="VENDAS_normal_sem_corresp_receber", index=False)
    ficha_financeira_vendas.to_excel(writer, sheet_name="VENDAS_ficha_financeira_vendas", index=False)
    df_status_incorreto.to_excel(writer, sheet_name="VENDAS_Dif_Status_Incorreta", index=False)
    mud_valor_tot_venda.to_excel(writer, sheet_name="VENDAS_mud_valor_vendas", index=False)
    mudancas_valor_venda.to_excel(writer, sheet_name="VENDAS_Aval_Mud_TotVenda", index=False)
    out_vendas_total.to_excel(writer, sheet_name="VENDAS_out_vendas_total", index=False)
    out_vendas_desconto.to_excel(writer, sheet_name="VENDAS_out_vendas_desconto", index=False)
    out_vendas_total_final.to_excel(writer, sheet_name="VENDAS_out_vendas_diff", index=False)
    vendas_cessao.to_excel(writer, sheet_name="VENDAS_vendas_cessao", index=False)

with pd.ExcelWriter("VENDAS_atualizações.xlsx") as writer:
    novas_vendas.to_excel(writer, sheet_name="Novas_vendas_mês", index=False)
    VENDAS_mod_campos.to_excel(writer, sheet_name="VENDAS_mod_campos", index=False)

#..............................................................................................................................................................................................
#..................................................................................CONTAS A RECEBER............................................................................................
#..............................................................................................................................................................................................

#Verificar se todas as parcelas correspondem a vendas com status “Normal”.
vendaID_Receber = BaseAtualizacao_Receber["vendaId"].isin(BaseAtualizacao_vendas[BaseAtualizacao_vendas["Status Venda"] == "NORMAL"]["vendaId"])
dados_verificacoes.append(["BASE CONTAS A RECEBER", "Verificar se todas as parcelas correspondem a vendas com status “Normal”.", vendaID_Receber.all()])
nao_corresp_vendaID_Receber = BaseAtualizacao_Receber[~vendaID_Receber]

#Agregar por ano-mês para fazer a próxima verificação. Para o período anterior ao ano-mês da atualização, a quantidade de parcelas da base atual tem que ser maior ou igual a de atualização. 
BaseAtualizacao_Receber["DataParcela_AnoMes"] = BaseAtualizacao_Receber["Data_Prc"].dt.strftime('%Y-%m')
BaseAtualizacao_Receber["DataParcela_AnoMes"] = pd.to_datetime(BaseAtualizacao_Receber["DataParcela_AnoMes"])

BaseAtual_Receber["DataParcela_AnoMes"] = BaseAtual_Receber["Data_Prc"].dt.strftime('%Y-%m')
BaseAtual_Receber["DataParcela_AnoMes"] = pd.to_datetime(BaseAtual_Receber["DataParcela_AnoMes"])
dt2 = datetime.strptime(dt, "%Y-%m")
#Agregando as informações por mês - Atual e Atualizada
receber_mensal_atualizacao = BaseAtualizacao_Receber[BaseAtualizacao_Receber["DataParcela_AnoMes"] < dt2].groupby("DataParcela_AnoMes").agg({"DataParcela_AnoMes": "count"})
receber_mensal_atualizacao.columns = ["Contagem_parcelas_atualizacao"]
receber_mensal_atualizacao.reset_index(inplace=True)
receber_mensal_atual = BaseAtual_Receber[BaseAtual_Receber["DataParcela_AnoMes"] < dt2].groupby("DataParcela_AnoMes").agg({"DataParcela_AnoMes": "count"})
receber_mensal_atual.columns = ["Contagem_parcelas_atual"]
receber_mensal_atual.reset_index(inplace=True)
#Juntando as tabelas e comparando os valores por periodo
mensal_atual_atualizacao = pd.merge(receber_mensal_atual, receber_mensal_atualizacao, on = "DataParcela_AnoMes", how = "outer" )
mensal_atual_atualizacao["Check_Parcelas"] = mensal_atual_atualizacao.Contagem_parcelas_atual >= mensal_atual_atualizacao.Contagem_parcelas_atualizacao
Mes_incorreto = mensal_atual_atualizacao[mensal_atual_atualizacao["Check_Parcelas"] == False]

if Mes_incorreto.shape[0] == 0:
    mes_erro = True
else:
    mes_erro = False
dados_verificacoes.append(["BASE CONTAS A RECEBER", "Verificação por mes-ano da quantidade de parcelas.", mes_erro])

#Verificar por Venda_Id, se a quantidade total de parcelas é igual ou menor.
#Agregando as informações das parcelas por vendaID para a base atualização
resumo_vendaId_atualizacao = BaseAtualizacao_Receber.groupby("vendaId").agg({"vendaId": "count", "Num Parcela": ["min", "max"], "Qtde Tot Parcela": "max"})
resumo_vendaId_atualizacao.columns = ["Contagem_Atualizacao", "Mim_num_parcela_Atualizacao", "Max_num_parcela_Atualizacao", "Qtde_Tot_Parcela"]
resumo_vendaId_atualizacao.reset_index(inplace=True)
#Agregando as informações das parcelas por vendaID para a base atual
resumo_vendaId_atual = BaseAtual_Receber.groupby("vendaId").agg({"vendaId": "count", "Num Parcela": "max"})
resumo_vendaId_atual.columns = ["Contagem_Atual", "Max_num_parcela_Atual"]
resumo_vendaId_atual.reset_index(inplace=True)
#Juntando as informações e fazendo as comparações
base_comparar_receber =  pd.merge(resumo_vendaId_atualizacao, resumo_vendaId_atual, on = "vendaId", how = "outer")
base_comparar_receber["Num_Parce_numeracao"] = base_comparar_receber.Max_num_parcela_Atualizacao - base_comparar_receber.Mim_num_parcela_Atualizacao + 1
base_comparar_receber["Check_num_parcela"] = base_comparar_receber.Contagem_Atualizacao == base_comparar_receber.Num_Parce_numeracao
base_comparar_receber["Check_contagem"] = base_comparar_receber.Contagem_Atualizacao <= base_comparar_receber.Contagem_Atual
base_comparar_receber["Reparcelamento_Atualizacao"] = base_comparar_receber.Max_num_parcela_Atualizacao > base_comparar_receber.Max_num_parcela_Atual
#Verificação da contagem de parcelas por venda ID. Comparando Atual x Atualização
contagem_receber_distinta = base_comparar_receber[(base_comparar_receber["Contagem_Atualizacao"] > 0) & (base_comparar_receber["Contagem_Atual"] > 0) & (base_comparar_receber["Check_contagem"] == False) & (base_comparar_receber["Reparcelamento_Atualizacao"] == False)]
if contagem_receber_distinta.shape[0] == 0:
    contagem_parcela = True
else:
    contagem_parcela = False
dados_verificacoes.append(["BASE CONTAS A RECEBER", "Comparação numero de parcelas atual x atualização.", contagem_parcela])

#Verificação da sequência de parcelas
sequencia_parcela_distinta = base_comparar_receber[(base_comparar_receber["Contagem_Atualizacao"] > 0) & (base_comparar_receber["Contagem_Atual"] > 0) & (base_comparar_receber["Check_num_parcela"] == False)]
if sequencia_parcela_distinta.shape[0] == 0:
    sequencia_parcela = True
else:
    sequencia_parcela = False
dados_verificacoes.append(["BASE CONTAS A RECEBER", "Verificação da sequência de parcelas.", sequencia_parcela])

#Verificação de Parcela.
BaseAtual_Receber["chave_verifica_parcela"] = BaseAtual_Receber["Descrição Parcela"].astype(str)+BaseAtual_Receber["Num Parcela"].astype(str) + BaseAtual_Receber["vendaId"].astype(str)+BaseAtual_Receber["Data_Prc"].astype(str)
BaseAtualizacao_Receber["chave_verifica_parcela"] = BaseAtualizacao_Receber["Descrição Parcela"].astype(str)+BaseAtualizacao_Receber["Num Parcela"].astype(str) + BaseAtualizacao_Receber["vendaId"].astype(str)+BaseAtualizacao_Receber["Data_Prc"].astype(str) 
BaseAtualizacao_Recebido["dataVencimento"] = BaseAtualizacao_Recebido["dataVencimento"].dt.strftime('%Y-%m-%d')
BaseAtualizacao_Recebido["chave_verifica_parcela"] = BaseAtualizacao_Recebido["Descrição Parcela"].astype(str)+BaseAtualizacao_Recebido["Num Parcela"].astype(str) + BaseAtualizacao_Recebido["vendaId"].astype(str)+BaseAtualizacao_Recebido["dataVencimento"].astype(str)
mascara_confere_parcela = BaseAtual_Receber["chave_verifica_parcela"].isin(BaseAtualizacao_Receber["chave_verifica_parcela"])
parcelas_nao_corresp = BaseAtual_Receber[~mascara_confere_parcela]
#Parcelas sem constar na base de atualização - Parcelas Pagas
mascara_parcela_paga = parcelas_nao_corresp["chave_verifica_parcela"].isin(BaseAtualizacao_Recebido["chave_verifica_parcela"])
parcelas_nao_corresp["Parcela_Paga"] = mascara_parcela_paga
#Parcelas sem constar na base de atualização - Vendas Cancelas
mascara_vendas_canceladas = parcelas_nao_corresp["vendaId"].isin(BaseAtualizacao_vendas[BaseAtualizacao_vendas["Status Venda"] == "CANCELADA"]["vendaId"])
parcelas_nao_corresp["Venda_Cancelada"] = mascara_vendas_canceladas
#Parcelas sem constar na base de atualização - Reparcelamentos
mascara_parcelas_reparceladas = parcelas_nao_corresp["vendaId"].isin(base_comparar_receber[base_comparar_receber["Reparcelamento_Atualizacao"] ==  True]["vendaId"])
parcelas_nao_corresp["Parcelas_Reparceladas"] = mascara_parcelas_reparceladas

#Verifica todas as mudancas de parcelas. Adicionada aqui porque precis ada chave de parcela
RECEBER_mod_campos = registros_mod(BaseAtual_Receber, BaseAtualizacao_Receber, "chave_verifica_parcela")

#Novas Parcelas
mascara_novas_percelas = BaseAtualizacao_Receber["chave_verifica_parcela"].isin(BaseAtual_Receber["chave_verifica_parcela"])
novas_parcelas_receber = BaseAtualizacao_Receber[~mascara_novas_percelas]

#Verificação Valor das parcelas que continuam na base para serem pagas
parcelas_continuam_atualizacao = BaseAtual_Receber[mascara_confere_parcela]
base_verificar_valor_parcela = parcelas_continuam_atualizacao[["Descrição Parcela", "Num Parcela", "Data_Prc", "Parcela Corrigida", "ValParcela_crc", "chave_verifica_parcela"]]
base_atualizacao_verificar_parcela = BaseAtualizacao_Receber[["chave_verifica_parcela", "ValParcela_crc"]]
verifica_valor_parcela = pd.merge(base_verificar_valor_parcela, base_atualizacao_verificar_parcela, on = "chave_verifica_parcela", how = "left")
verifica_valor_parcela["Valor_OK"] = verifica_valor_parcela.ValParcela_crc_x <= verifica_valor_parcela.ValParcela_crc_y
parcelas_valor_verificar = verifica_valor_parcela[verifica_valor_parcela["Valor_OK"] == False]

if parcelas_valor_verificar.shape[0] == 0:
    valor_parcela = True
else:
    valor_parcela = False
dados_verificacoes.append(["BASE CONTAS A RECEBER", "Verificação valor parcelas atual x atualizacao.", valor_parcela])

#Mudancas no valor principal da parcela
mudancas_valor_parcela = RECEBER_mod_campos[RECEBER_mod_campos['Valor_Prc'] != ""]

#Salvando as informacoes verificadas no Excel
with pd.ExcelWriter("RECEBER_verificacoes.xlsx") as writer:
    nao_corresp_vendaID_Receber.to_excel(writer, sheet_name="RECEBER_nao_corresp_vendaID_Receber", index=False)
    Mes_incorreto.to_excel(writer, sheet_name="RECEBER_Mes_incorreto", index=False)
    contagem_receber_distinta.to_excel(writer, sheet_name="RECEBER_contagem_receber_distinta", index=False)
    sequencia_parcela_distinta.to_excel(writer, sheet_name="RECEBER_sequencia_parcela_distinta", index=False)
    parcelas_nao_corresp.to_excel(writer, sheet_name="RECEBER_Parcelas_Out_atualizacao", index=False)
    parcelas_valor_verificar.to_excel(writer, sheet_name="RECEBER_parcelas_valor_verificar", index=False)
    mudancas_valor_parcela.to_excel(writer, sheet_name="mudancas_valor_parcela", index=False)

with pd.ExcelWriter("RECEBER_atualizacoes.xlsx") as writer:
    RECEBER_mod_campos.to_excel(writer, sheet_name="RECEBER_mod_campos", index=False)
    novas_parcelas_receber.to_excel(writer, sheet_name="RECEBER_novas_parcelas", index=False)

#..............................................................................................................................................................................................
#..................................................................................CONTAS RECEBIDAS............................................................................................
#..............................................................................................................................................................................................
#Verificar se a quantidade única de VendaId é menor ou igual, comparando atual com atualizada.
vendas_id_atual = BaseAtual_Recebido.vendaId.unique().size
vendas_id_atualizacao = BaseAtualizacao_Recebido.vendaId.unique().size
check_qtd_vendaId = vendas_id_atual <= vendas_id_atualizacao
dados_verificacoes.append(["BASE CONTAS RECEBIDAS", "Verificar se a quantidade única de VendaId é menor ou igual atual x atualizacao.", check_qtd_vendaId])

#Verificar se a quantidade e valor total das parcelas, por vendaId é menor ou igual, atual com atualizada. 
resumo_parcelaspagas_atual = BaseAtual_Recebido.groupby("vendaId").agg({"vendaId": "count", "TotParcel": "sum"})
resumo_parcelaspagas_atual.columns = ["parcelas_pagas_atual", "total_recebido_atual"]
resumo_parcelaspagas_atual.reset_index(inplace=True)
resumo_parcelaspagas_atualizacao = BaseAtualizacao_Recebido.groupby("vendaId").agg({"vendaId": "count", "TotParcel": "sum"})
resumo_parcelaspagas_atualizacao.columns = ["parcelas_pagas_atualizacao", "total_recebido_atualizacao"]
resumo_parcelaspagas_atualizacao.reset_index(inplace=True)
verifica_valor_parcela_paga = pd.merge(resumo_parcelaspagas_atual, resumo_parcelaspagas_atualizacao, on = "vendaId", how = "left")

#Faz o check de quantidade de parcelas pagas e total recebido
verifica_valor_parcela_paga["Check_num_parcelas"] = verifica_valor_parcela_paga.parcelas_pagas_atual <= verifica_valor_parcela_paga.parcelas_pagas_atualizacao
verifica_valor_parcela_paga["Check_total_recebido"] = (verifica_valor_parcela_paga.total_recebido_atual - verifica_valor_parcela_paga.total_recebido_atualizacao) <= 2

#Verificar se todas as parcelas recebidas são de vendasId que constam na tabela vendas
mascara_parcelas_vendas = BaseAtualizacao_Recebido.vendaId.isin(BaseAtualizacao_vendas["vendaId"])
dados_verificacoes.append(["BASE CONTAS RECEBIDAS", "Verificar se todas as parcelas recebidas são de vendasId que constam na tabela vendas.", mascara_parcelas_vendas.all()])
nao_corresp_base_vendas = BaseAtualizacao_Recebido[~mascara_parcelas_vendas]

#Verificação de Parcela. Fazer a chave “DescParcela+Num Parcela+VendaID+DataVencimento” na base atual e atualização. Verificar a correspondência Base Atual na Base Atualização. Verificar também o campo “TotalParcel”, tem que permanecer o mesmo. 
BaseAtual_Recebido["dataVencimento"] = BaseAtual_Recebido["dataVencimento"].dt.strftime('%Y-%m-%d')
BaseAtual_Recebido["chave_verifica_parcela"] = BaseAtual_Recebido["Descrição Parcela"].astype(str)+BaseAtual_Recebido["Num Parcela"].astype(str) + BaseAtual_Recebido["vendaId"].astype(str)+BaseAtual_Recebido["dataVencimento"].astype(str)
baseAtual_parcelas_pagas_valor =  BaseAtual_Recebido[["chave_verifica_parcela", "dataVencimento", "TotParcel"]]
baseAtual_parcelas_pagas_valor.columns = ["chave_verifica_parcela", "dataVencimento_Atual", "TotParcel_Atual"]
baseAtual_parcelas_pagas_valor = baseAtual_parcelas_pagas_valor.groupby(["chave_verifica_parcela", "dataVencimento_Atual"]).sum()
baseAtual_parcelas_pagas_valor = baseAtual_parcelas_pagas_valor.reset_index()
baseAtualizacao_parcelas_pagas_valor =  BaseAtualizacao_Recebido[["chave_verifica_parcela", "dataVencimento", "TotParcel"]]
baseAtualizacao_parcelas_pagas_valor.columns = ["chave_verifica_parcela", "dataVencimento_Atualizacao", "TotParcel_Atualizacao"]
baseAtualizacao_parcelas_pagas_valor = baseAtualizacao_parcelas_pagas_valor.groupby(["chave_verifica_parcela", "dataVencimento_Atualizacao"]).agg({"TotParcel_Atualizacao": 'sum', "chave_verifica_parcela": 'count'})
baseAtualizacao_parcelas_pagas_valor.columns = ['TotParcel_Atualizacao', 'Total_pagamentos_parcela']
baseAtualizacao_parcelas_pagas_valor = baseAtualizacao_parcelas_pagas_valor.reset_index()
parcelas_inconsistentes = pd.merge(baseAtual_parcelas_pagas_valor, baseAtualizacao_parcelas_pagas_valor, on = "chave_verifica_parcela", how = "left")
parcelas_inconsistentes["Check_Data"] = parcelas_inconsistentes['dataVencimento_Atual'] == parcelas_inconsistentes['dataVencimento_Atualizacao']
parcelas_inconsistentes["Check_ValorPago"] = (parcelas_inconsistentes['TotParcel_Atual'] - parcelas_inconsistentes['TotParcel_Atualizacao']) > 1
parcelas_inconsistentes["Paga_Parcialmente"] = parcelas_inconsistentes['Total_pagamentos_parcela'] > 1

#Para as parcelas pagas no mês de atualização fazer a verificação do valor pago, “TotParcel+TotDesc”, com o valor que deveria ser pago, base atual – contas a receber - ValParcela_crc, considerando uma margem.
parcelas_pagas_mes = BaseAtualizacao_Recebido[(BaseAtualizacao_Recebido["anoRecebimento"] == dt2.year) & (BaseAtualizacao_Recebido["mesRecebimento"] == dt2.month)]
parcelas_pagas_referencia = BaseAtualizacao_Recebido[(BaseAtualizacao_Recebido["anoRecebimento"] == dt2.year) & (BaseAtualizacao_Recebido["mesRecebimento"] == dt2.month)]
parcelas_pagas_referencia["Total_receber"] = parcelas_pagas_referencia["TotParcel"] + parcelas_pagas_referencia["TotDesc"]
base_parcelas = BaseAtual_Receber[["ValParcela_crc", "chave_verifica_parcela"]]
parcelas_pagas_referencia = pd.merge(parcelas_pagas_referencia, base_parcelas, on = "chave_verifica_parcela", how = "left")
parcelas_pagas_referencia["Dif_pago_previsto"] = parcelas_pagas_referencia["Total_receber"] - parcelas_pagas_referencia["ValParcela_crc"] 
parcelas_pagas_referencia["Check_Valor_Pago_Distoa"] =  parcelas_pagas_referencia["Dif_pago_previsto"].abs() > 50

#Verificar as mudanças nos campos
BaseAtual_Recebido["chave_verifica_parcela_recebida"] = BaseAtual_Recebido["Descrição Parcela"].astype(str)+BaseAtual_Recebido["numParcGer"].astype(str) + BaseAtual_Recebido["vendaId"].astype(str)+BaseAtual_Recebido["Data Recebimento"].astype(str)
BaseAtualizacao_Recebido["chave_verifica_parcela_recebida"] = BaseAtualizacao_Recebido["Descrição Parcela"].astype(str)+BaseAtualizacao_Recebido["numParcGer"].astype(str) + BaseAtualizacao_Recebido["vendaId"].astype(str)+BaseAtualizacao_Recebido["Data Recebimento"].astype(str)
#Corresp atual na atualização
mascara_corresp_parcelas_pagas = BaseAtual_Recebido["chave_verifica_parcela"].isin(BaseAtualizacao_Recebido["chave_verifica_parcela"])
nao_corresp_parcelas_pagas = BaseAtual_Recebido[~mascara_corresp_parcelas_pagas]
#Corresp atualização na atual
mascara_corresp_parcelas_pagas_atualizacao = BaseAtualizacao_Recebido["chave_verifica_parcela_recebida"].isin(BaseAtual_Recebido["chave_verifica_parcela_recebida"])
parcelas_pagas_por_dif = BaseAtualizacao_Recebido[~mascara_corresp_parcelas_pagas_atualizacao]
RECEBIDO_mod_campos = registros_mod(BaseAtual_Recebido, BaseAtualizacao_Recebido, "chave_verifica_parcela_recebida")

#Salvando as Mudancas no Excel
with pd.ExcelWriter("RECEBIDO_verificacoes.xlsx") as writer:
    verifica_valor_parcela_paga.to_excel(writer, sheet_name="RECEBIDO_verifica_num_valor_parcela_paga", index=False)
    nao_corresp_base_vendas.to_excel(writer, sheet_name="RECEBIDO_nao_corresp_base_vendas", index=False)
    parcelas_inconsistentes.to_excel(writer, sheet_name="RECEBIDO_parcelas_inconsistentes", index=False)
    parcelas_pagas_referencia.to_excel(writer, sheet_name="RECEBIDO_parcelas_pagas_referencia", index=False)

with pd.ExcelWriter("RECEBIDO_atualizacoes.xlsx") as writer:
    parcelas_pagas_mes.to_excel(writer, sheet_name="RECEBIDO_parcelas_pagas_mes", index=False)
    RECEBIDO_mod_campos.to_excel(writer, sheet_name="RECEBIDO_mod_campos", index=False)
    parcelas_pagas_por_dif.to_excel(writer, sheet_name="RECEBIDO_parcelas_pagas_mes_diff", index=False)

#Fazendo o documento word com as verificações

gerar_relatorio(dados_verificacoes)
